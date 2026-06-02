from __future__ import annotations

from datetime import date, datetime, timedelta
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Pt

from .config import DB_PATH, REPORT_DIR
from .database import connect, init_db


def query_articles(start_date: str, end_date: str, db_path: Path | str = DB_PATH) -> pd.DataFrame:
    init_db(db_path)
    with connect(db_path) as con:
        df = pd.read_sql_query(
            """
            SELECT
                article_id,
                first_seen_date,
                publication_date,
                journal_name,
                title,
                authors,
                doi,
                url,
                fulltext_url,
                source,
                topics,
                matched_keywords,
                study_type,
                status,
                favorite,
                user_notes,
                personal_tags,
                abstract
            FROM articles
            WHERE first_seen_date BETWEEN ? AND ?
            ORDER BY first_seen_date DESC, journal_name, title
            """,
            con,
            params=(start_date, end_date),
        )
    return df


def _explode_topics(df: pd.DataFrame) -> pd.DataFrame:
    rows = []
    if df.empty:
        return pd.DataFrame(columns=["topic", "title", "journal_name", "first_seen_date"])
    for _, row in df.iterrows():
        for topic in str(row.get("topics", "")).split(";"):
            topic = topic.strip()
            if topic:
                rows.append({
                    "topic": topic,
                    "title": row.get("title", ""),
                    "journal_name": row.get("journal_name", ""),
                    "first_seen_date": row.get("first_seen_date", ""),
                })
    return pd.DataFrame(rows)


def export_excel(df: pd.DataFrame, output_path: Path) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    with pd.ExcelWriter(output_path, engine="openpyxl") as writer:
        public_cols = [
            "first_seen_date", "publication_date", "journal_name", "title", "authors", "doi",
            "url", "fulltext_url", "source", "topics", "matched_keywords", "study_type",
            "status", "favorite", "personal_tags", "user_notes", "abstract"
        ]
        df[[c for c in public_cols if c in df.columns]].to_excel(writer, index=False, sheet_name="articles")

        summary = (
            df.groupby(["journal_name"], dropna=False)
            .size()
            .reset_index(name="new_articles")
            .sort_values("new_articles", ascending=False)
        ) if not df.empty else pd.DataFrame(columns=["journal_name", "new_articles"])
        summary.to_excel(writer, index=False, sheet_name="journal_summary")

        topics = _explode_topics(df)
        topic_summary = topics.groupby("topic").size().reset_index(name="articles").sort_values("articles", ascending=False) if not topics.empty else pd.DataFrame(columns=["topic", "articles"])
        topic_summary.to_excel(writer, index=False, sheet_name="topic_summary")
        topics.to_excel(writer, index=False, sheet_name="topic_hits")

        reading = df[df.get("favorite", 0).astype(str).isin(["1", "True", "true"])] if not df.empty and "favorite" in df.columns else pd.DataFrame()
        reading.to_excel(writer, index=False, sheet_name="favorites")
    return output_path


def _add_table_from_df(doc: Document, df: pd.DataFrame, columns: list[str], max_rows: int = 50) -> None:
    if df.empty:
        doc.add_paragraph("无记录。")
        return
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    for i, col in enumerate(columns):
        table.rows[0].cells[i].text = col
    for _, row in df.head(max_rows).iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(columns):
            cells[i].text = str(row.get(col, ""))[:1200]


def export_word_report(df: pd.DataFrame, output_path: Path, title: str, start_date: str, end_date: str) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"].font.size = Pt(10.5)

    doc.add_heading(title, level=0)
    doc.add_paragraph(f"统计周期：{start_date} 至 {end_date}")
    doc.add_paragraph(f"新增论文总数：{len(df)} 篇")

    if df.empty:
        doc.add_paragraph("本周期未发现新增论文，或抓取源未返回新记录。")
        doc.save(output_path)
        return output_path

    topic_df = _explode_topics(df)
    topic_hit_n = int(df["topics"].fillna("").astype(str).str.len().gt(0).sum()) if "topics" in df.columns else 0
    fav_n = int((df.get("favorite", 0) == 1).sum()) if "favorite" in df.columns else 0
    doc.add_paragraph(f"专题命中论文：{topic_hit_n} 篇；收藏论文：{fav_n} 篇。")

    doc.add_heading("一、本周期期刊更新概览", level=1)
    journal_summary = df.groupby("journal_name").size().reset_index(name="新增篇数").sort_values("新增篇数", ascending=False)
    _add_table_from_df(doc, journal_summary.rename(columns={"journal_name": "期刊"}), ["期刊", "新增篇数"], max_rows=40)

    doc.add_heading("二、专题命中概览", level=1)
    if topic_df.empty:
        doc.add_paragraph("本周期未命中专题词库。")
    else:
        topic_summary = topic_df.groupby("topic").size().reset_index(name="命中篇数").sort_values("命中篇数", ascending=False)
        _add_table_from_df(doc, topic_summary.rename(columns={"topic": "专题"}), ["专题", "命中篇数"], max_rows=30)

    doc.add_heading("三、专题命中论文", level=1)
    topic_articles = df[df["topics"].fillna("").astype(str) != ""].copy() if "topics" in df.columns else pd.DataFrame()
    if topic_articles.empty:
        doc.add_paragraph("本周期无专题命中论文。")
    else:
        topic_articles = topic_articles.rename(columns={"journal_name": "期刊", "title": "题名", "topics": "专题标签", "doi": "DOI", "study_type": "研究类型"})
        _add_table_from_df(doc, topic_articles, ["期刊", "题名", "专题标签", "研究类型", "DOI"], max_rows=60)

    doc.add_heading("四、收藏 / 精读池论文", level=1)
    fav_df = df[(df.get("favorite", 0) == 1) | (df.get("status", "") == "精读")].copy()
    if fav_df.empty:
        doc.add_paragraph("本周期无收藏或精读状态论文。")
    else:
        fav_df = fav_df.rename(columns={"journal_name": "期刊", "title": "题名", "status": "状态", "user_notes": "备注", "doi": "DOI"})
        _add_table_from_df(doc, fav_df, ["期刊", "题名", "状态", "备注", "DOI"], max_rows=50)

    doc.add_heading("五、全部新增论文", level=1)
    all_df = df.rename(columns={"journal_name": "期刊", "title": "题名", "topics": "专题标签", "status": "阅读状态", "doi": "DOI"})
    _add_table_from_df(doc, all_df, ["期刊", "题名", "专题标签", "阅读状态", "DOI"], max_rows=250)

    doc.save(output_path)
    return output_path


def make_daily_report(day: str | None = None, db_path: Path | str = DB_PATH) -> tuple[Path, Path]:
    day = day or date.today().isoformat()
    df = query_articles(day, day, db_path)
    xlsx = REPORT_DIR / "daily" / f"journal_tracker_daily_{day}.xlsx"
    docx = REPORT_DIR / "daily" / f"journal_tracker_daily_{day}.docx"
    export_excel(df, xlsx)
    export_word_report(df, docx, f"体育科学期刊更新日报 {day}", day, day)
    return xlsx, docx


def make_weekly_report(end_day: str | None = None, db_path: Path | str = DB_PATH) -> tuple[Path, Path]:
    end = datetime.strptime(end_day, "%Y-%m-%d").date() if end_day else date.today()
    start = end - timedelta(days=6)
    start_s, end_s = start.isoformat(), end.isoformat()
    df = query_articles(start_s, end_s, db_path)
    iso_year, iso_week, _ = end.isocalendar()
    xlsx = REPORT_DIR / "weekly" / f"journal_tracker_weekly_{iso_year}-W{iso_week:02d}.xlsx"
    docx = REPORT_DIR / "weekly" / f"journal_tracker_weekly_{iso_year}-W{iso_week:02d}.docx"
    export_excel(df, xlsx)
    export_word_report(df, docx, f"体育科学期刊周度文献情报简报 {iso_year}-W{iso_week:02d}", start_s, end_s)
    return xlsx, docx
